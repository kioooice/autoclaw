# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm

doc = Document(r'C:\Users\Administrator\Desktop\碱锌挂镀五彩钝化及孔内防锈.docx')

print('='*70)
print('【完整格式问题检查报告】')
print('='*70)

# 1. 页面设置问题
print()
print('一、页面设置')
print('-'*50)
for i, section in enumerate(doc.sections):
    print(f'上边距: {section.top_margin.cm:.2f} cm (标准: 2.5cm)')
    print(f'下边距: {section.bottom_margin.cm:.2f} cm (标准: 2.5cm)')
    print(f'左边距: {section.left_margin.cm:.2f} cm (标准: 3.0cm)')
    print(f'右边距: {section.right_margin.cm:.2f} cm (标准: 2.5cm)')

# 2. 标题格式检查
print()
print('二、标题格式问题')
print('-'*50)

for para in doc.paragraphs:
    if para.style and 'Heading' in para.style.name:
        for run in para.runs:
            if run.text.strip():
                font_size = run.font.size.pt if run.font.size else '继承样式'
                bold = run.font.bold if run.font.bold is not None else '继承'
                print(f'{para.style.name}: "{para.text[:30]}" - 字号={font_size}pt, 加粗={bold}')
                break

# 3. 检查是否有图片
print()
print('三、图片检查')
print('-'*50)
inline_shapes = doc.inline_shapes
print(f'文档中图片数量: {len(inline_shapes)}')

# 4. 检查参考文献
print()
print('四、参考文献检查')
print('-'*50)
has_refs = False
for para in doc.paragraphs:
    if '参考文献' in para.text or 'References' in para.text:
        has_refs = True
        print(f'找到参考文献部分')
        break
if not has_refs:
    print('未找到参考文献部分')

# 5. 检查目录
print()
print('五、目录检查')
print('-'*50)
has_toc = False
for para in doc.paragraphs:
    if '目录' in para.text or '目 录' in para.text:
        has_toc = True
        print(f'找到目录: {para.text}')
        break
if not has_toc:
    print('未找到目录部分')

# 6. 检查摘要
print()
print('六、摘要检查')
print('-'*50)
has_abstract = False
for para in doc.paragraphs:
    if '摘要' in para.text or 'Abstract' in para.text:
        has_abstract = True
        print(f'找到摘要: {para.text}')
        break
if not has_abstract:
    print('未找到摘要部分')

# 7. 检查关键词
print()
print('七、关键词检查')
print('-'*50)
has_keywords = False
for para in doc.paragraphs:
    if '关键词' in para.text or 'Keywords' in para.text:
        has_keywords = True
        print(f'找到关键词: {para.text[:50]}')
        break
if not has_keywords:
    print('未找到关键词部分')

# 8. 检查章节编号
print()
print('八、章节编号检查')
print('-'*50)
for para in doc.paragraphs:
    if para.style and para.style.name == 'Heading 1':
        text = para.text
        if text and not any(char.isdigit() for char in text[:3]):
            print(f'一级标题缺少编号: "{text}"')

# 9. 表格标题格式
print()
print('九、表格标题格式检查')
print('-'*50)
for para in doc.paragraphs:
    if para.style and para.style.name == '表标题':
        print(f'表格标题: "{para.text[:40]}"')

# 10. 统计
print()
print('='*70)
print('【格式问题总结】')
print('='*70)